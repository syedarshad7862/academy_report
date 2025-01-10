from openpyxl import load_workbook
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# for docx to pdf
from docx2pdf import convert
# Output folder for saving docx files
output_folder = r'C:\Users\ThinkPad\Desktop\python project\\'

# Load the workbook
wb = load_workbook(filename = r'C:\Users\ThinkPad\Desktop\academy report\master_sheet.xlsx', data_only=True)

# it is for rating sheet
rating_sheet = wb["Rating"]

skip_row_of_rating = 3
start_cell_of_rating = 2

rating_sheet_data = {}
sheet_data_attendance = []
sheet_data_assignment = []
sheet_data_mock = []
# rating_names = ["Attendance report"]
for i, row in enumerate(rating_sheet.rows, start=1):
      if i <= skip_row_of_rating:
            continue
      print(f"row {i} row_data {row[start_cell_of_rating].value}")
      row_data = []
      for j, cell in enumerate(row[start_cell_of_rating:], start=start_cell_of_rating):
            print(cell.value)
            row_data.append(cell.value)
            # rating_sheet_data["data"] = row_data
      if row_data:
            if len(sheet_data_attendance) < 3:
                 sheet_data_attendance.append(row_data)
            elif len(sheet_data_assignment) < 3:
                 sheet_data_assignment.append(row_data)
            else:
                 sheet_data_mock.append(row_data)           
rating_sheet_data["attendance report"] = sheet_data_attendance
rating_sheet_data["assignment report"] = sheet_data_assignment
rating_sheet_data["mock report"] = sheet_data_mock


attendance_rating_criteria = rating_sheet_data["attendance report"]
assignment_rating_criteria = rating_sheet_data["assignment report"]
mock_test_rating_criteria = rating_sheet_data["mock report"]

# function for percentage of attendance rating
def get_rating(p):
    for criteria in attendance_rating_criteria:
        attendance_report, rating, lower_bound, upper_bound, action, empty, empty2,areas_of_improvement, final_score_basis,overall_feedback= criteria
        # print(rating)
        if  lower_bound <= p <= upper_bound:
            return attendance_report, rating, action,areas_of_improvement, final_score_basis
    return "unknown", "no action"

# function for percentage of assignment rating
def get_assignment_rating(p):
    for criteria in assignment_rating_criteria: #i have change here
        assignment_report, rating, lower_bound, upper_bound, action, deadline, empty, areas_of_improvement, final_score_basis, empty2 = criteria
        # print(rating)
        if  lower_bound <= p <= upper_bound:
            return assignment_report, rating, action, deadline,areas_of_improvement,final_score_basis
    return "unknown", "no action", "no deadline"

# function for percentage of mock test rating
def get_mock_test_rating(p):
    for criteria in mock_test_rating_criteria:
        mock_test_report, rating, lower_bound, upper_bound, action, deadline, mock_test_feedback, areas_of_improvement, final_score_basis, empty = criteria
        # print(rating)
        if  lower_bound <= p <= upper_bound:
            return mock_test_report,rating, action, deadline, mock_test_feedback,areas_of_improvement,final_score_basis
    return "unknown", "no action", "no deadline"
# function for feedback of mock test
def get_feed_back(p):
    for criteria in mock_test_rating_criteria:
        _, rating, lower_bound, upper_bound, action, deadline, mock_test_feedback,areas_of_improvement, final_score_basis, empty = criteria
        # print(rating)
        if  lower_bound <= p <= upper_bound:
            return mock_test_feedback
    return "no feedback"

# function for overall report and final socre
def get_final_score(p):
    for criteria in attendance_rating_criteria:
        attendance_report, rating, lower_bound, upper_bound, action, empty, empty2,areas_of_improvement, final_score_basis, overall_feedback= criteria
        # print(rating)
        if  lower_bound <= p <= upper_bound:
            return overall_feedback
    return "unknown", "no action"
# Extract the valid date ranges (start_date, end_date) from the 'report gen' sheet
# it is for report gen/ taking the dates
report_gen = wb["report gen"]
skip_row_of_report_gen = 3
start_cell_of_report_gen = 2
dates = []
for i, row in enumerate(report_gen.iter_rows(min_row=skip_row_of_report_gen + 1), start=skip_row_of_report_gen + 1):
    # Extract start_date and end_date from specific columns
    start_date = row[start_cell_of_report_gen].value  # Assuming start_date is in this column
    end_date = row[start_cell_of_report_gen + 1].value  # Assuming end_date is in the next column

    # Ensure the values are valid before processing
    if isinstance(start_date, datetime) and isinstance(end_date, datetime):
        print(f"Row {i}: Start Date = {start_date}, End Date = {end_date}")
        dates.append((start_date,end_date))
    else:
        print(f"Row {i}: Invalid date values - Start Date: {start_date}, End Date: {end_date}")

# save the dates in seperate variables
if dates:
    start_date, end_date = dates[0]
    
# Sheet names and other configuration
# sheet_names = ["Attendance report"]
# sheet_names = ["Attendance report", "Assignment report", "Mock Test Report", "Mock Test Meta"]

# # Define the start_date and end_date
# start_date = datetime.strptime("02-12-2024", "%d-%m-%Y")
# end_date = datetime.strptime("09-12-2024", "%d-%m-%Y")

# Dictionary to store the  of attendance sheet
attendance_data = {}
# Loop through all sheets
# for sheet_name in sheet_names:
sheet = wb["Attendance report"]
for row in sheet:
    # print(row)
    is_session_row = False
    start_saving = False
    is_name_row = False
    sheet_data = []
    row_data = []
    name_row = []
    for col in row:
        row_data.append(col.value)
        if col.value == "Session Date":
            is_session_row = True
            
        elif col.value == "Session Number":
            pass
        elif col.value  is None:
            pass
        else:
            is_name_row = True
        # if col.value == "Assignment Date":
        #     is_assignment_date = True 
        # elif col.value == "Assignment Name":
        #     is_assignment_names = True
        # elif col.value == "Due":
        #     pass
        # elif col.value == "Student Name":
        #     pass
        # elif   col.value == "Submitted":
        #     pass    
    r_val = []
    if is_name_row or isinstance(row_data, datetime):
        # check none value
        name_row.append([data for data in row_data if data is not None])
        attendance_data[name_row[0][0]] = name_row[0][1:]
        
        
    # if is_session_row:
    #     print(row_data)
    #     for i, r in enumerate(row_data):
    #         if isinstance(r, datetime):
    #             if r >= start_date and r <= end_date:
    #                 r_val.append(r.strftime("%d-%m-%Y"))
                        # sheet_data.append(row_data)
                    # elif r > end_date:
                    #     start_saving = False
                    #     break        
        # if is_assignment_date:
        #     print("it is assignment row")
        #     assignment_row.append([data for data in row_data if data is not None])
        #     assignment_data[name_row[0][0]] = assignment_row[0][1:]
        # if is_assignment_names:
        #     print("it is assignment name row")    
        
print(attendance_data)  
start_index = 0
end_index = 0
attendance_range_data = {}
for key,value in attendance_data.items():
    print(key)
    if key == "Session Date":
        for index, date_temp in enumerate(value):
            if start_date == date_temp:
                start_index = index
            if end_date == date_temp:
                end_index = index    
    elif key == "Session Number":
        pass
    else:
        actual_data = value[start_index: end_index +1]
        print(actual_data)
        print(key,actual_data) 
        attendance_range_data[key] = actual_data
        print("after actual data")      
    print(value)
    print("after value")
print(attendance_range_data)

assignment_data = {}
sheet = wb["Assignment report"]
for row in sheet:
    # print(row)
    is_assignment_row = False
    is_assignment_names = False
    is_assignment_date = False
    is_student_name = False
    row_data = []
    name_row = []
    for col in row:
        row_data.append(col.value)
        if col.value == "Assignment Date":
            is_assignment_date = True 
        elif col.value == "Assignment Name":
            is_assignment_names = True
        elif col.value is None:
            pass
        elif col.value is None:
            pass
        elif   col.value is None:
            pass 
        else:
            is_student_name = True   
    r_val = []
    if is_student_name:
        # check none value
        name_row.append([data for data in row_data if data is not None])
        assignment_data[name_row[0][0]] = name_row[0][1:]
        print(f"inside assignment loop {name_row}")
        
        
    # if is_assignment_date:
    #     print(row_data)
    #     for i, r in enumerate(row_data):
    #         if isinstance(r, datetime):
    #             if r >= start_date and r <= end_date:
    #                 r_val.append(r.strftime("%d-%m-%Y"))
                        # sheet_data.append(row_data)
                    # elif r > end_date:
                    #     start_saving = False
                    #     break        
        # if is_assignment_date:
        #     print("it is assignment row")
        #     assignment_row.append([data for data in row_data if data is not None])
        #     assignment_data[name_row[0][0]] = assignment_row[0][1:]
        # if is_assignment_names:
        #     print("it is assignment name row")    
        
print(assignment_data) 

start_index = 0
end_index = 0
assignment_range_data = {}
for key,value in assignment_data.items():
    if key == 'Assignment Date':
        for index, date_temp in enumerate(value):
            if start_date == date_temp:
                start_index = index
            if end_date >= date_temp:
                end_index = index
    elif key == "Student Name":
        pass
    else:
        actual_data = value[start_index:end_index +1]
        print(actual_data)    
        print(key,actual_data) 
        assignment_range_data[key] = actual_data
        print("after actual data")

mock_test_data = {}
sheet = wb["Mock Test Report"]
for row in sheet:
    # print(row)
    is_mock_test_row = False
    is_mock_test_names = False
    is_mock_test_date = False
    is_sd_name = False
    row_data = []
    name_row = []
    for col in row:
        row_data.append(col.value)
        if col.value == "Mock Test Date":
            is_mock_test_date = True 
        elif col.value == "Mock Test Name":
            is_mock_test_names = True
        elif col.value is None:
            pass
        elif col.value == "Student Name":
            pass
        elif   col.value is None:
            pass 
        else:
            is_sd_name = True   
    r_val = []
    if is_sd_name:
        # check none value
        name_row.append([data for data in row_data if data is not None])
        mock_test_data[name_row[0][0]] = name_row[0][1:]
        print(f"inside assignment loop {name_row}")    
        
print(mock_test_data)

m_start_index = 0
m_end_index = 0
mock_test_range_data = {}
for key,value in mock_test_data.items():
    if key == "Mock Test Date":
        for index, date_temp in enumerate(value):
            if start_date == date_temp:
                m_start_index = index
            if end_date >= date_temp:
                m_end_index = index
    elif key == "Student Name":
        pass
    else:
        actual_data = value[m_start_index:m_end_index + 1]
        mock_test_range_data[key] = actual_data
print(mock_test_range_data)

# mock test Meta sheet
sheet = wb['Mock Test Meta']
mock_test_meta_sheet_data = []
for i,row in enumerate(sheet,start=1):
    if i <= 3:
        continue
    is_grade_row = False
    is_grades_row = False
    row_data = []
    name_row = []
    for col in row[2:]:
        row_data.append(col.value)
    if row_data:
                
        mock_test_meta_sheet_data.append(row_data)

# make dictionary of mock_test_meta for grades
grade_action_mapping = {row[0] : {"action": row[1], "score": row[2]} for row in mock_test_meta_sheet_data if row[0] is not None }           
#function for assignment names table start here
def parse_assignment_report(data):
    
    # Extract student data
    student_data = []
    for key,value in data.items():
        if key == "Assignment Name":
            pass
        else:
            student = {
                "name": key,
                "assignments": value, # Assignment submission status
                # "total_due": len(submitted)
                }
            student_data.append(student)
    
    return student_data  
# function for pending assignments
def get_pending_assignments(assignments, student):
    pending = []
    for assignment, submitted in zip(assignments, student["assignments"]):
        if submitted == 0:  # Check for unsubmitted (0)
            pending.append({
                "name": assignment,
                # "rating": student["rating"],
                # "action_needed": student["action_needed"],
            })
            
    return pending 
def parse_mock_test_report(data):
    # Extract assignment names
    # mock_test_names = data[1][1:6]  # Get the columns for mock test names
    
    # Extract student data
    student_data = []
    for key,value in data.items():
        if key == 'Mock Test Name':  # Check if the key is equal to mock test name
            pass
        else:
            student = {
                "name": key,
                "mock_test_grade": value,  # mock test submission status
                # "total_due": row[6],
                # "total_submitted": row[7],
            }
            student_data.append(student)
    
    return student_data
def get_moct_test(mock_test_names, student):
    pending = []
    for mock_test, submitted in zip(mock_test_names, student["mock_test_grade"]):
        if submitted == 0 or submitted == "A" or submitted == "B" or submitted == "C" or submitted == "D" or submitted == "E":  # Check for unsubmitted (0)
            pending.append({
                "name": mock_test,
                "grade": student["mock_test_grade"],
                # "action_needed": student["action_needed"],
            })
            
    return pending
s = 0
# Loop through each student row and generate a report
for key,value in attendance_range_data.items():
    if key == "Session Number":
        continue
    student_name =  key
    attended = [present for present in value if present != 0]
    print(attended, len(attended))
    attended_sessions = len(attended)   # Attended sessions count
    total_sessions = len(value)
    attendance_percentage = (attended_sessions/total_sessions) * 100 # Attendance percentage
    print(f"attended: {attended_sessions} and {attendance_percentage}")
    
    # calling function for attendance rating
    # attendance_report,rating, action_needed,areas_of_improvement, final_score_basis = get_rating(attendance_percentage)
    # print(rating,action_needed)
    
        # Create a new Document
    document = Document()

    title1 = document.add_paragraph('Daari Academy – Your Path to Success')
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title = document.add_paragraph("Student Performance Report")
    sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_title = document.add_paragraph("Course name: CMA US - Part 2")
    course_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_name = document.add_paragraph("Report Name: Weekly Report")
    report_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_duration = document.add_paragraph()
    report_duration.add_run(f"Report Duration: From {start_date.strftime('%d-%m-%Y')} to {end_date.strftime('%d-%m-%Y')}")
    report_duration.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line = report_duration._element
    p_pr = line.get_or_add_pPr()
    p_borders = parse_xml(
        r'<w:pBdr {}>'
        r'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/>'
        r'</w:pBdr>'.format(nsdecls('w'))
    )
    p_pr.append(p_borders)
    
    paragraph = document.add_paragraph(f"Student Name: ")
    paragraph.add_run(student_name).font.color.rgb = RGBColor(0, 0, 255)  # RGB for blue
    line2 = paragraph._element
    p_pr = line.get_or_add_pPr()
    p_borders = parse_xml(
        r'<w:pBdr {}>'
        r'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/>'
        r'</w:pBdr>'.format(nsdecls('w'))
    )
    p_pr.append(p_borders)
    
    attendance_title = document.add_paragraph()
    attendance_title.add_run("A. Attendance Report")
    attendance_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Attendance report table
    table = document.add_table(rows=1, cols=3)

    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Total Live Session Scheduled'
    hdr_cells[1].text = 'Attended'
    hdr_cells[2].text = 'Percentage'
     
    # table.style = "Grid Table 6 Colorfull - Accent 5"
    table.style = "Light Grid"
        # convert the 1 to 100%
    percentage = int(attendance_percentage)
        # Add attendance data
    row_cells = table.add_row().cells
    row_cells[0].text = str(total_sessions)  # Total scheduled sessions
    row_cells[1].text = str(attended_sessions)
    row_cells[2].text = str(f"{percentage}%")
    print(percentage)
        # calling function for attendance rating
    attendance_report,rating, action_needed,areas_of_improvement, final_score_basis = get_rating(attendance_percentage)
        #Rating section
    paragraph2 = document.add_paragraph(f"Rating: ")
    paragraph2.add_run(rating).font.color.rgb = RGBColor(0, 0, 255)
    
    paragraph3 = document.add_paragraph(f"Action needed:")
    paragraph3.add_run(action_needed).font.color.rgb = RGBColor(0,0,255)
    
    #Assignment report start here
    asssignment_title = document.add_paragraph()
    asssignment_title.add_run("B. Assignment Report")
    asssignment_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    std_assignment = assignment_range_data[student_name]
    submitted = [assignment for assignment in std_assignment if assignment != 0]
    assignment_name = assignment_range_data['Assignment Name']
    assignment_percentage = (len(submitted)/len(assignment_name)) * 100 # assignment percentage
    
            # Assignment report table
    table = document.add_table(rows=1, cols=3)
    table.style = "Light Grid"
    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Assignment Due'
    hdr_cells[1].text = 'Submitted'
    hdr_cells[2].text = 'Percentage'
    
    # convert the 0.0 to 0
    converted_assignment_percentage = int(assignment_percentage)
    # Add assignment data
    row_cells = table.add_row().cells
    row_cells[0].text = str(len(assignment_name))  # Total assignment
    row_cells[1].text = str(len(submitted))
    row_cells[2].text = str(f"{converted_assignment_percentage}%")
        # calling function for assignment rating
    assignment_report,assignment_rating, assignment_action_needed, deadline,areas_of_improvement_in_assignment,assignment_score = get_assignment_rating(assignment_percentage)
    print(assignment_rating,assignment_action_needed)
    
    #Rating section
    asm_rating = document.add_paragraph(f"Rating: ")
    asm_rating.add_run(assignment_rating).font.color.rgb = RGBColor(0, 0, 255)  # RGB for blue
    act_needed = document.add_paragraph(f"Action needed: ")
    act_needed.add_run(assignment_action_needed).font.color.rgb = RGBColor(0, 0, 255)  # RGB for blue
    
    # Assignments overdue table
    # assignments_overdue = document.add_paragraph("Assignments overdue")
    # assignments_overdue.alignment = WD_ALIGN_PARAGRAPH.CENTER
    student_data = parse_assignment_report(assignment_range_data)
    test_data = student_data[s]
    
    selected_student_name = student_name # student_name
    selected_student = test_data["name"]

    if selected_student:
        pending_assignments = get_pending_assignments(assignment_name, test_data)
        print(f"Pending assignments for {selected_student_name}:")
        if pending_assignments:
                # assignment detail table
            # Add header row
            document.add_paragraph("Assignments overdue").alignment = WD_ALIGN_PARAGRAPH.CENTER
            table3 = document.add_table(rows=1, cols=2)
            hdr_cells = table3.rows[0].cells
            hdr_cells[0].text = 'Assignment/s are to be submitted.'
            hdr_cells[1].text = 'Deadline'
            table3.style = "Light Grid"
        for n, assignment in enumerate(pending_assignments, start=1) :
            # print(f"- {assignment['name']} | Rating: {assignment['rating']} | Action: {assignment['action_needed']}")
            print(f"- {assignment['name']},Deadline: {deadline}")


            
            row_cells = table3.add_row().cells
            row_cells[0].text = str(assignment["name"])
            row_cells[1].text = str(deadline)
        
    else:
        print(f"No data found for student: {selected_student_name}")
        
    #Mock test report start here
    attendance_title = document.add_paragraph()
    attendance_title.add_run("C. Mock Test Report")
    attendance_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    std_mock_test = mock_test_range_data['Mock Test Name']
    mock_test_due = len(std_mock_test)
    mock_test_attended = mock_test_range_data[student_name]
    attended = [mock_attended for mock_attended in mock_test_attended if mock_attended != 0]
    mock_test_percentage = (len(attended)/mock_test_due) * 100 # mock test percentage

       # mock test report table
    table3 = document.add_table(rows=1, cols=3)
    table3.style = "Light Grid"
    
        # Add header row
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'Mock Test Conducted'
    hdr_cells[1].text = 'Mock Test Attended'
    hdr_cells[2].text = 'Percentage'
    
    # convert the 0.0 to 0
    converted_moct_test_percentage = int(mock_test_percentage)
    # Add mock test data
    row_cells = table3.add_row().cells
    row_cells[0].text = str(mock_test_due)  # Total mock test
    row_cells[1].text = str(len(attended))
    row_cells[2].text = str(f"{converted_moct_test_percentage}%")
    
        # calling function for mock test rating
    mock_test_report,mock_test_rating, mock_test_action_needed, mock_deadline, mock_test_feedback,areas_of_improvement_in_mock,mock_test_score = get_mock_test_rating(mock_test_percentage)
    
        # calling function for mock test rating
    mock_test_report,mock_test_rating, mock_test_action_needed, mock_deadline, mock_test_feedback,areas_of_improvement_in_mock,mock_test_score = get_mock_test_rating(mock_test_percentage)
    # print(mock_test_rating,mock_test_action_needed, mock_test_feedback)
    #Rating section
    mock_rating = document.add_paragraph(f"Rating: ")
    mock_rating.add_run(f"{mock_test_rating}").font.color.rgb = RGBColor(0,0,255)
    mock_action = document.add_paragraph(f"Action needed: ")
    mock_action.add_run(f"{mock_test_action_needed}").font.color.rgb = RGBColor(0,0,255) 
    
    student_of_mock_test = parse_mock_test_report(mock_test_range_data)
    mock_data = student_of_mock_test[s]
    search_student_name = student_name # student_name
    mock_selected_student = mock_data["name"]
    result = []
    if mock_selected_student:
        mock_names = get_moct_test(std_mock_test, mock_data)
        grades = mock_data["mock_test_grade"]
        print(f"mock test  for {search_student_name}:")
        student_result = {"name": student_name, "grades": []}
        for mock_test, grade in zip(mock_names, grades):
            if grade in grade_action_mapping:
                print(grade)
                action_details = grade_action_mapping[grade]
                student_result["grades"].append({
                 "mock_test_name": mock_test,   
                "grade": grade,
                "action": action_details["action"],
                "score": action_details["score"]
            })
            else:
                 student_result["grades"].append({
                "grade": grade,
                "action": "No action available",
                "score": None
                })
        result.append(student_result)         
    else:
        print(f"No data found for student: {search_student_name}")
    total_score = 0    
    table4 = document.add_table(rows=1, cols=3)
    table4.style = "Light Grid"
        # Add header row
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = 'Mock Test Name'
    hdr_cells[1].text = 'Grade'
    hdr_cells[2].text = 'Future Action'
    for student in result:
        print(f"Student Name: {student['name']}")
        for grade_detail in student["grades"]:
            total_score += grade_detail['score']
            # Add mock test data
            row_cells = table4.add_row().cells
            row_cells[0].text = str(grade_detail['mock_test_name']['name'])  # Total mock test
            row_cells[1].text = str(grade_detail['grade'])
            row_cells[2].text = str(grade_detail['action'])
            # print(f"mock test name: {grade_detail['mock_test_name']['name']}  Grade: {grade_detail['grade']} | Action: {grade_detail['action']} | Score: {grade_detail['score']} | total_score = {total_score}")
        print("-" * 50) 
    print(f"{(total_score/50)*100}") 
    mock_percentage = (total_score/50)*100
    Overall_Score_rating = document.add_paragraph("Overall Score rating: ")
    Overall_Score_rating.add_run(f" Based on the Grade & assigned score (A-10, B-8, C-6, D-4, E-2) = {int(mock_percentage)}%").font.color.rgb = RGBColor(0, 0, 255) 
    feed_back = get_feed_back(mock_percentage)
    print(feed_back)   
    mock_test_feedback = document.add_paragraph(f"FeedBack:")
    mock_test_feedback.add_run(feed_back).font.color.rgb = RGBColor(0, 0, 255)
    
    # overall report start here
    overall_report = document.add_paragraph("Overall Report")
    overall_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
     #Overall Report
    table5 = document.add_table(rows=1, cols=3)
    table5.style = "Light Grid"
        # Add header row
    hdr_cells = table5.rows[0].cells
    hdr_cells[0].text = 'Title'
    hdr_cells[1].text = 'Ratin/ Overall Grade'
    hdr_cells[2].text = 'Area of Improvement'
    
    row1 = table5.add_row().cells
    row1[0].text = str(attendance_report)
    row1[1].text = str(rating)
    row1[2].text = str(areas_of_improvement)
    row2 = table5.add_row().cells
    row2[0].text = str(assignment_report)
    row2[1].text = str(assignment_rating)
    row2[2].text = str(areas_of_improvement_in_assignment)
    row3 = table5.add_row().cells
    row3[0].text = str(mock_test_report)
    row3[1].text = str(mock_test_rating)
    row3[2].text = str(areas_of_improvement_in_mock)
    final_score = final_score_basis + mock_test_score + assignment_score
    # print(final_score/30*100)
    final_score_percentage = (final_score/30)*100
    overall_feedback = get_final_score(final_score_percentage)
    # print(overall_feedback)
    final_score = document.add_paragraph("Final Score: ")
    final_score.add_run(f"{int(final_score_percentage)}%").font.color.rgb = RGBColor(0,0,225)
    feedback = document.add_paragraph("Feedback: ")
    feedback.add_run(overall_feedback).font.color.rgb = RGBColor(0,0,255)
    file_name = f"{student_name.replace(' ', '_')}.docx"
    output_path = f"{output_folder}{file_name}"
    s = s + 1
    document.save(output_path)
    # convert
    convert(output_path)
    print(f"Document saved for {student_name}: {output_path}")
        